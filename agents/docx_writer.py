"""
Agent 6 — docx_writer.py
Builds a PhD-format Word document from research/proposal_draft.json and
research/case_study_data.json, saving to output/proposal_ifrs_union_bank.docx.
"""

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Style constants ────────────────────────────────────────────────────────────
HEADING_COLOR = RGBColor(0x1A, 0x37, 0x6C)  # navy blue
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
REF_SIZE = Pt(11)
TITLE_SIZE = Pt(16)
CH_SIZE = Pt(14)    # chapter heading
SEC_SIZE = Pt(12)   # sub-section (1.1, 2.1 …)
SUB_SIZE = Pt(11)   # sub-sub-section (2.1.1, 2.2.1 …)
LINE_SPACE = 1.5
PARA_AFTER = Pt(12)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_font(run, size=None, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    if size:
        run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _apply_line_spacing(p):
    """Apply 1.5 line spacing and 12pt space-after to a paragraph."""
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = LINE_SPACE
    p.paragraph_format.space_after = PARA_AFTER


def _chapter_heading(doc: Document, text: str) -> None:
    """Level 1 heading — chapter, references, appendices."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_font(run, size=CH_SIZE, bold=True, color=HEADING_COLOR)


def _section_heading(doc: Document, text: str) -> None:
    """Level 2 heading — numbered sections (1.1, 2.1 …)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_font(run, size=SEC_SIZE, bold=True, color=HEADING_COLOR)


def _sub_heading(doc: Document, text: str) -> None:
    """Level 3 heading — numbered sub-sections (2.1.1, 2.2.1 …)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, size=SUB_SIZE, bold=True, italic=True, color=HEADING_COLOR)


def _body(doc: Document, text: str) -> None:
    """Add a body paragraph with 1.5 line spacing and 12pt after."""
    if not text:
        return
    p = doc.add_paragraph(style="Normal")
    _apply_line_spacing(p)
    run = p.add_run(str(text))
    _set_font(run)


def _numbered_list(doc: Document, items: list) -> None:
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"{i}. {item}")
        _set_font(run)


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def _table_header_row(table, headers: list[str]) -> None:
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        _set_font(run, bold=True)
    _shade_row(row, hex_color="DAEEF3")


def _add_table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        _set_font(run, size=Pt(10))


def _shade_row(row, hex_color: str = "D9E1F2") -> None:
    """Apply background shading to a table row."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def _render_gaps(doc: Document, gaps) -> None:
    """Render research gap items with bold 'GAP N:' prefix."""
    if not isinstance(gaps, list):
        _body(doc, str(gaps))
        return
    for gap in gaps:
        gap_str = str(gap).strip()
        p = doc.add_paragraph(style="Normal")
        _apply_line_spacing(p)
        colon_idx = gap_str.find(":") if gap_str.upper().startswith("GAP") else -1
        if colon_idx != -1:
            bold_run = p.add_run(gap_str[: colon_idx + 1])
            _set_font(bold_run, bold=True)
            rest_run = p.add_run(gap_str[colon_idx + 1:])
            _set_font(rest_run)
        else:
            run = p.add_run(gap_str)
            _set_font(run)


def _render_hypotheses(doc: Document, hypotheses) -> None:
    """Render H0N / HNA labelled hypothesis pairs."""
    if not isinstance(hypotheses, list):
        _body(doc, str(hypotheses))
        return
    for i, hyp in enumerate(hypotheses, start=1):
        if isinstance(hyp, dict):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.space_after = Pt(2)
            h0 = p.add_run(f"H0{i}: ")
            _set_font(h0, bold=True)
            null_run = p.add_run(str(hyp.get("null", "")))
            _set_font(null_run)
            p2 = doc.add_paragraph(style="Normal")
            p2.paragraph_format.space_after = Pt(8)
            ha = p2.add_run(f"H{i}A: ")
            _set_font(ha, bold=True)
            alt = hyp.get("alternative") or hyp.get("alternate", "")
            alt_run = p2.add_run(str(alt))
            _set_font(alt_run)
        else:
            _body(doc, str(hyp))


def _render_terms(doc: Document, terms) -> None:
    """Render operational definition of terms."""
    if isinstance(terms, dict):
        for term, definition in terms.items():
            p = doc.add_paragraph(style="Normal")
            _apply_line_spacing(p)
            bold_run = p.add_run(f"{term}: ")
            _set_font(bold_run, bold=True)
            def_run = p.add_run(str(definition).strip())
            _set_font(def_run)
    elif terms:
        _body(doc, str(terms))


def _render_var_table(doc: Document, variables) -> None:
    """Render the variable operationalisation table."""
    if not (isinstance(variables, list) and variables):
        _body(doc, str(variables))
        return
    headers = ["Variable", "Proxy Measure", "Data Source", "Time Period"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _table_header_row(table, headers)
    for var in variables:
        if isinstance(var, dict):
            _add_table_row(table, [
                var.get("variable", ""), var.get("proxy", ""),
                var.get("source", ""), var.get("period", ""),
            ])
        else:
            _add_table_row(table, [str(var), "", "", ""])
    doc.add_paragraph()


# ── Main builder ───────────────────────────────────────────────────────────────

def build_document(proposal: dict, case_study: dict) -> Document:
    doc = Document()
    _set_margins(doc)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE

    # ── Nested chapter accessors with flat-key fallback ───────────────────────
    ch1 = proposal.get("chapter_one") or {}
    ch2 = proposal.get("chapter_two") or {}
    ch3 = proposal.get("chapter_three") or {}
    _old_lit  = proposal.get("literature_review") or {}
    _old_meth = proposal.get("methodology") or {}

    _bg    = ch1.get("background_to_study") or proposal.get("introduction", "")
    _sop   = ch1.get("statement_of_problem") or proposal.get("statement_of_problem", "")
    _obj   = ch1.get("research_objectives") or proposal.get("research_objectives", [])
    _qs    = ch1.get("research_questions") or proposal.get("research_questions", [])
    _hyp   = ch1.get("research_hypotheses") or proposal.get("research_hypotheses", [])
    _sig   = ch1.get("significance_of_study") or proposal.get("significance_of_study", "")
    _scope = ch1.get("scope_and_delimitation") or proposal.get("scope_and_delimitation", "")
    _terms = ch1.get("operational_definition_of_terms") or {}

    # ── 1. Cover page ──────────────────────────────────────────────────────────
    cover = proposal.get("cover_page", {})
    if isinstance(cover, dict):
        title_text       = cover.get("title", "")
        author_text      = cover.get("author", "")
        institution_text = cover.get("institution", "")
        date_text        = cover.get("date", datetime.now().strftime("%B %Y"))
    else:
        title_text = str(cover)
        author_text = institution_text = ""
        date_text = datetime.now().strftime("%B %Y")

    for _ in range(4):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(
        title_text or (
            "The Effect of International Financial Reporting Standards (IFRS) on "
            "Going Concern Assessment and Financial Stability: "
            "A Case Study of Union Bank of Nigeria PLC (2012\u20132022)"
        )
    )
    _set_font(title_run, size=TITLE_SIZE, bold=True, color=HEADING_COLOR)

    for line in [author_text, institution_text, date_text]:
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            _set_font(run)

    doc.add_page_break()

    # ── Table of Contents placeholder ─────────────────────────────────────────
    _chapter_heading(doc, "TABLE OF CONTENTS")
    p = doc.add_paragraph()
    run = p.add_run(
        "[Please update the Table of Contents in Microsoft Word: "
        "References \u2192 Update Table \u2192 Update entire table]"
    )
    run.font.italic = True
    run.font.name = FONT_NAME
    doc.add_page_break()

    # ── CHAPTER ONE: INTRODUCTION ─────────────────────────────────────────────
    _chapter_heading(doc, "CHAPTER ONE \u2014 INTRODUCTION")

    _section_heading(doc, "1.1 Background to the Study")
    _body(doc, str(_bg))

    _section_heading(doc, "1.2 Statement of the Problem")
    _body(doc, str(_sop))

    _section_heading(doc, "1.3 Research Objectives")
    if isinstance(_obj, list):
        _numbered_list(doc, _obj)
    else:
        _body(doc, str(_obj))

    _section_heading(doc, "1.4 Research Questions")
    if isinstance(_qs, list):
        _numbered_list(doc, _qs)
    else:
        _body(doc, str(_qs))

    _section_heading(doc, "1.5 Research Hypotheses")
    _render_hypotheses(doc, _hyp)

    _section_heading(doc, "1.6 Significance of the Study")
    _body(doc, str(_sig))

    _section_heading(doc, "1.7 Scope and Delimitation of the Study")
    _body(doc, str(_scope))

    _section_heading(doc, "1.8 Operational Definition of Terms")
    _render_terms(doc, _terms)

    doc.add_page_break()

    # ── CHAPTER TWO: LITERATURE REVIEW ───────────────────────────────────────
    _chapter_heading(doc, "CHAPTER TWO \u2014 LITERATURE REVIEW")

    # 2.1 Conceptual Review
    _section_heading(doc, "2.1 Conceptual Review")
    conceptual = ch2.get("conceptual_review") or {}
    if isinstance(conceptual, dict) and conceptual:
        _sub_heading(doc, "2.1.1 IFRS: Concept, Adoption and Framework")
        _body(doc, str(conceptual.get("ifrs_concept_adoption", "")))
        _sub_heading(doc, "2.1.2 Going Concern: Concept and Regulatory Basis")
        _body(doc, str(conceptual.get("going_concern_concept", "")))
        _sub_heading(doc, "2.1.3 Banking Sector Stability: Measures and Determinants")
        _body(doc, str(conceptual.get("banking_stability", "")))
        _sub_heading(doc, "2.1.4 Expected Credit Loss (ECL) under IFRS 9")
        _body(doc, str(conceptual.get("expected_credit_loss_ifrs9", "")))
    else:
        _body(doc, str(conceptual or _old_lit.get("conceptual_review", "")))

    # 2.2 Theoretical Framework
    _section_heading(doc, "2.2 Theoretical Framework")
    theo = ch2.get("theoretical_framework") or {}
    if isinstance(theo, dict) and theo:
        _sub_heading(doc, "2.2.1 Agency Theory")
        _body(doc, str(theo.get("agency_theory", "")))
        _sub_heading(doc, "2.2.2 Signalling Theory")
        _body(doc, str(theo.get("signalling_theory", "")))
        _sub_heading(doc, "2.2.3 Stakeholder Theory")
        _body(doc, str(theo.get("stakeholder_theory", "")))
        _sub_heading(doc, "2.2.4 Positive Accounting Theory")
        _body(doc, str(theo.get("positive_accounting_theory", "")))
    else:
        _body(doc, str(theo
                       or _old_lit.get("theoretical_review", "")
                       or proposal.get("theoretical_framework", "")))

    # 2.3 Empirical Review
    _section_heading(doc, "2.3 Empirical Review")
    empirical = ch2.get("empirical_review") or {}
    if isinstance(empirical, dict) and empirical:
        _sub_heading(doc, "2.3.1 IFRS and Banking Stability \u2014 International Evidence")
        _body(doc, str(empirical.get("ifrs_banking_stability_international", "")))
        _sub_heading(doc, "2.3.2 IFRS and Going Concern \u2014 African Evidence")
        _body(doc, str(empirical.get("ifrs_going_concern_africa", "")))
        _sub_heading(doc, "2.3.3 IFRS 9 ECL and Credit Risk in Nigerian Banks")
        _body(doc, str(empirical.get("ifrs9_ecl_nigerian_banks", "")))
        _sub_heading(doc,
                     "2.3.4 Union Bank of Nigeria PLC \u2014 Prior Studies and Context")
        _body(doc, str(empirical.get("union_bank_prior_studies", "")))
    else:
        _body(doc, str(empirical or _old_lit.get("empirical_review", "")))

    # 2.4 Research Gaps
    _section_heading(doc, "2.4 Identification of Research Gaps")
    gaps = ch2.get("research_gaps") or proposal.get("research_gaps", [])
    _render_gaps(doc, gaps)

    # 2.5 Summary of Literature Review
    _section_heading(doc, "2.5 Summary of Literature Review")
    _body(doc, str(ch2.get("summary_of_literature", "")))

    doc.add_page_break()

    # ── CHAPTER THREE: RESEARCH METHODOLOGY ──────────────────────────────────
    _chapter_heading(doc, "CHAPTER THREE \u2014 RESEARCH METHODOLOGY")

    _section_heading(doc, "3.1 Research Design")
    _body(doc, str(ch3.get("research_design") or _old_meth.get("research_design", "")))

    _section_heading(doc, "3.2 Research Philosophy")
    _body(doc, str(
        ch3.get("research_philosophy") or _old_meth.get("research_philosophy", "")))

    _section_heading(doc, "3.3 Population and Sample")
    _body(doc, str(
        ch3.get("population_and_sample")
        or _old_meth.get("population_and_sampling", "")))

    _section_heading(doc, "3.4 Sources of Data")
    _body(doc, str(
        ch3.get("sources_of_data") or _old_meth.get("data_collection", "")))

    _section_heading(doc, "3.5 Data Collection Procedure")
    _body(doc, str(ch3.get("data_collection_procedure", "")))

    _section_heading(doc, "3.6 Operationalisation of Variables")
    variables = (ch3.get("variable_operationalisation")
                 or proposal.get("data_sources_and_variables", []))
    _render_var_table(doc, variables)

    _section_heading(doc, "3.7 Model Specification")
    model_text = str(
        ch3.get("model_specification") or _old_meth.get("model_specification", ""))
    # Render equation on its own line; definitions indented beneath it
    where_m = re.search(r"\bwhere\b[:\s]", model_text, re.IGNORECASE)
    if where_m:
        _body(doc, model_text[: where_m.start()].strip())
        for defn in re.split(r"[;\n]", model_text[where_m.start():]):
            defn = defn.strip()
            if defn:
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(defn)
                _set_font(run)
    else:
        lines = [ln.strip() for ln in model_text.split("\n") if ln.strip()]
        for idx, line in enumerate(lines):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.space_after = Pt(3)
            if idx > 0:
                p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line)
            _set_font(run)

    _section_heading(doc, "3.8 Method of Data Analysis")
    _body(doc, str(
        ch3.get("method_of_data_analysis")
        or _old_meth.get("analytical_techniques", "")))

    _section_heading(doc, "3.9 Validity and Reliability of Secondary Data")
    _body(doc, str(ch3.get("validity_and_reliability", "")))

    _section_heading(doc, "3.10 Ethical Considerations")
    _body(doc, str(
        ch3.get("ethical_considerations")
        or proposal.get("ethical_considerations", "")))

    doc.add_page_break()

    # ── REFERENCES ─────────────────────────────────────────────────────────────
    _chapter_heading(doc, "REFERENCES")
    references = proposal.get("references", [])
    if isinstance(references, list):
        for ref in references:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(str(ref))
            run.font.name = FONT_NAME
            run.font.size = REF_SIZE
    else:
        _body(doc, str(references))

    doc.add_page_break()

    # ── APPENDIX A ─────────────────────────────────────────────────────────────
    _chapter_heading(doc,
                     "APPENDIX A \u2014 Union Bank of Nigeria PLC "
                     "Financial Data Table (2015\u20132022)")

    financial_data = case_study.get("financial_data_by_year", {})
    if financial_data:
        app_headers = ["Year", "NPL (%)", "CAR (%)", "ROA (%)",
                       "GC Modified", "Audit Opinion"]
        app_table = doc.add_table(rows=1, cols=len(app_headers))
        app_table.style = "Table Grid"
        _table_header_row(app_table, app_headers)
        for year, metrics in sorted(financial_data.items()):
            _add_table_row(
                app_table,
                [
                    year,
                    str(metrics.get("NPL_pct", "")),
                    str(metrics.get("CAR_pct", "")),
                    str(metrics.get("ROA_pct", "")),
                    "Yes" if metrics.get("going_concern_modified") else "No",
                    metrics.get("audit_opinion", ""),
                ],
            )
        doc.add_paragraph()

    source_note = doc.add_paragraph()
    note_run = source_note.add_run(
        "Sources: Union Bank of Nigeria PLC Annual Reports (2015\u20132022); "
        "CBN Banking Supervision Reports; NGX regulatory filings. "
        "GC Modified = Going Concern paragraph or Emphasis of Matter issued by auditors."
    )
    note_run.font.italic = True
    note_run.font.name = FONT_NAME
    note_run.font.size = Pt(10)

    doc.add_page_break()

    # ── APPENDIX B ─────────────────────────────────────────────────────────────
    _chapter_heading(doc, "APPENDIX B \u2014 List of Variables, Proxies and Sources")

    appendices = proposal.get("appendices", {})
    appendix_b = ""
    if isinstance(appendices, dict):
        appendix_b = appendices.get("appendix_b_variables", "")
    _body(doc,
          str(appendix_b) if appendix_b else
          "See Section 3.6 (Operationalisation of Variables) in Chapter Three "
          "for the complete variable operationalisation table.")

    return doc


def main() -> None:
    proposal_path = Path("research/proposal_draft.json")
    case_study_path = Path("research/case_study_data.json")

    if not proposal_path.exists():
        raise FileNotFoundError(f"{proposal_path} not found — run synthesis_agent first")
    if not case_study_path.exists():
        raise FileNotFoundError(f"{case_study_path} not found — run case_study_agent first")

    proposal = json.loads(proposal_path.read_text())
    case_study = json.loads(case_study_path.read_text())

    doc = build_document(proposal, case_study)

    output_dir = Path("output")
    output_dir.mkdir(exist_ok=True)

    output_path = output_dir / "proposal_ifrs_union_bank.docx"
    doc.save(str(output_path))
    print(f"[docx_writer] Document saved to {output_path}")


if __name__ == "__main__":
    main()

