"""
Agent 5 — synthesis_agent.py
Calls OpenRouter free-tier LLM to synthesise all research into a PhD proposal.
Uses requests only — NO anthropic SDK.
Research sources: Tavily search/extract/deep-research results and Union Bank case study data.
"""

import json
import os
import re
import sys
from pathlib import Path

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

OPENROUTER_BASE = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL = "openrouter/free"

MAX_CONTEXT_CHARS = 28_000  # ~8k tokens safety threshold


def load_research() -> dict:
    """Load and combine all research artefacts from Tavily and the case study."""
    research_dir = Path("research")

    tavily_path = research_dir / "tavily_results.json"
    tavily_research_path = research_dir / "tavily_research.json"
    case_study_path = research_dir / "case_study_data.json"
    queries_path = research_dir / "queries.json"

    tavily_results = []
    if tavily_path.exists():
        tavily_results = json.loads(tavily_path.read_text())[:60]

    tavily_research = []
    if tavily_research_path.exists():
        tavily_research = json.loads(tavily_research_path.read_text())

    case_study_data = {}
    if case_study_path.exists():
        case_study_data = json.loads(case_study_path.read_text())

    queries_data = {}
    if queries_path.exists():
        queries_data = json.loads(queries_path.read_text())

    return {
        "tavily_results": tavily_results,
        "tavily_research": tavily_research,
        "case_study_data": case_study_data,
        "queries_data": queries_data,
    }


def build_context(research: dict) -> str:
    """
    Build a combined context string from all three sources, prioritised by
    information density. The case study data is always placed last (highest
    priority anchor) and never truncated.
    """
    parts: list[str] = []

    # ── Source A: Tavily standard search + extract results ────────────────────
    for i, item in enumerate(research.get("tavily_results", []), start=1):
        tier = item.get("tier", "search")
        content_snippet = (
            item.get("raw_content") or item.get("content") or ""
        )[:500]
        answer = item.get("answer_snippet", "")
        parts.append(
            f"[TAV-{i}:{tier.upper()}] {item.get('title', '')} / "
            f"{item.get('source', '')} / "
            f"CLUSTER:{item.get('cluster', '')} / "
            f"{content_snippet}"
            + (f" / ANSWER: {answer[:200]}" if answer else "")
        )

    # ── Source B: Tavily deep research reports ────────────────────────────────
    for i, report in enumerate(research.get("tavily_research", []), start=1):
        cluster = report.get("cluster", "")
        # The report may come back as "report" string or nested "data"
        report_text = (
            report.get("report")
            or report.get("answer")
            or report.get("summary")
            or ""
        )
        sources = report.get("sources") or report.get("results") or []
        source_urls = " | ".join(
            s.get("url", s) if isinstance(s, dict) else str(s)
            for s in sources[:5]
        )
        parts.append(
            f"[DEEP-{i}] CLUSTER:{cluster} / "
            f"{report_text[:1500]}"
            + (f" / SOURCES: {source_urls}" if source_urls else "")
        )

    # ── Source C: Live-enriched case study (always last, never truncated) ─────
    case_study = research.get("case_study_data", {})
    # Exclude bulk live_web_snippets from the main JSON dump to save tokens;
    # they are already summarised in TAV results above.
    case_study_for_context = {
        k: v for k, v in case_study.items() if k != "live_web_snippets"
    }
    case_study_json = json.dumps(case_study_for_context, indent=1)

    # Live snippets summary (answer lines only, capped)
    live_snippets = case_study.get("live_web_snippets", [])
    live_summary_parts = []
    for s in live_snippets[:8]:
        ans = s.get("answer", "")
        url = s.get("url", "")
        content = (s.get("content", "") or "")[:300]
        if ans:
            live_summary_parts.append(f"LIVE[{url}]: {ans[:250]}")
        elif content:
            live_summary_parts.append(f"LIVE[{url}]: {content}")
    live_summary = "\n".join(live_summary_parts)

    case_study_section = (
        "\n--- UNION BANK LIVE + BASELINE CASE STUDY DATA ---\n"
        + (f"LIVE WEB EVIDENCE:\n{live_summary}\n\n" if live_summary else "")
        + f"STRUCTURED DATA:\n{case_study_json}"
    )
    parts.append(case_study_section)

    context = "\n\n".join(parts)

    # Truncate A/B/C from the END while always preserving case study section
    if len(context) > MAX_CONTEXT_CHARS:
        budget = MAX_CONTEXT_CHARS - len(case_study_section) - 200
        truncated_parts = []
        current_len = 0
        for part in parts[:-1]:
            if current_len + len(part) + 2 > budget:
                truncated_parts.append("[... context truncated to fit token limit ...]")
                break
            truncated_parts.append(part)
            current_len += len(part) + 2
        truncated_parts.append(case_study_section)
        context = "\n\n".join(truncated_parts)

    return context


def call_openrouter(
    system_prompt: str,
    user_message: str,
    max_tokens: int = 8000,
) -> str:
    """POST to OpenRouter and return the assistant message content."""
    api_key = os.environ["OPENROUTER_API_KEY"]
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://github.com/dukemawex/IFRX",
        "X-Title": "IFRS-PhD-Proposal-Generator",
    }

    payload = {
        "model": OPENROUTER_MODEL,
        "max_tokens": max_tokens,
        "stream": False,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
    }

    try:
        response = requests.post(OPENROUTER_BASE, headers=headers, json=payload, timeout=120)

        if not response.ok:
            raise RuntimeError(
                f"OpenRouter API error {response.status_code}: {response.text[:500]}"
            )

        data = response.json()

        if "error" in data:
            raise RuntimeError(
                f"OpenRouter returned an error in the response body "
                f"(HTTP {response.status_code}): {data['error']}"
            )

        choices = data.get("choices")
        if not choices:
            raise RuntimeError(
                f"OpenRouter response contained no choices "
                f"(HTTP {response.status_code}). Full response: {data}"
            )

        content = choices[0].get("message", {}).get("content")
        if not content:
            raise RuntimeError(
                f"OpenRouter choice[0] has no 'content' field "
                f"(HTTP {response.status_code}). Choice: {choices[0]}"
            )

        return content
    except RuntimeError:
        raise
    except Exception as exc:
        raise RuntimeError(f"OpenRouter request failed: {exc}") from exc


# ── Dissertation formatting ───────────────────────────────────────────────────

_HEADING_COLOR = RGBColor(0x1A, 0x37, 0x6C)  # navy blue
_FONT_NAME = "Times New Roman"
_BODY_PT = Pt(12)
_TITLE_PT = Pt(16)
_HEAD1_PT = Pt(14)
_HEAD2_PT = Pt(13)
_HEAD3_PT = Pt(12)


def _run(para, text: str, size=None, bold: bool = False, italic: bool = False,
         color=None):
    """Add a run with consistent font settings."""
    run = para.add_run(text)
    run.font.name = _FONT_NAME
    run.font.size = size or _BODY_PT
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _chapter_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    _run(p, text, size=Pt(14), bold=True, color=_HEADING_COLOR)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, size=Pt(12), bold=True, color=_HEADING_COLOR)


def _sub_heading(doc: Document, text: str) -> None:
    """Sub-subsection label (e.g. 2.1.1)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _run(p, text, size=Pt(11), bold=True, italic=True, color=_HEADING_COLOR)


def _body(doc: Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    _run(p, _enforce_ub_scope(str(text).strip()))


def _table_header(table, headers: list[str]) -> None:
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True)


def _table_row(table, values: list[str]) -> None:
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        _run(p, str(v), size=Pt(11))


def _split_paragraphs(text: str) -> list[str]:
    """Split text on one or more blank lines into non-empty paragraphs."""
    return [p.strip() for p in re.split(r"\n{2,}", text.strip()) if p.strip()]


# ── Scope-enforcement substitutions ──────────────────────────────────────────
# Replace any multi-bank / panel-study framing that may survive in the JSON
# with Union-Bank-only language, keeping the study as a single-entity case study.

_SCOPE_PATTERNS: list[tuple[re.Pattern, str]] = [
    # explicit "N listed Nigerian banks" and close variants
    (re.compile(r"\b\d+\s+listed\s+(?:Nigerian\s+)?banks?\b", re.IGNORECASE),
     "Union Bank of Nigeria PLC"),
    # "listed deposit money banks" (NDIC/CBN terminology for the broader sector)
    (re.compile(r"\blisted\s+deposit\s+money\s+banks?\b", re.IGNORECASE),
     "Union Bank of Nigeria PLC"),
    # "panel data across/of/from/for multiple banks" → longitudinal framing
    (re.compile(
        r"\bpanel\s+data\s+(?:across|of|from|for)\s+(?:multiple\s+)?(?:listed\s+)?banks?\b",
        re.IGNORECASE),
     "longitudinal data for Union Bank of Nigeria PLC"),
    # residual "multiple banks"
    (re.compile(r"\bmultiple\s+banks?\b", re.IGNORECASE),
     "Union Bank of Nigeria PLC"),
    # sector-wide / panel study framing
    (re.compile(r"\bsector-?wide\s+panel\s+study\b", re.IGNORECASE),
     "single-entity longitudinal case study"),
    (re.compile(r"\bpanel\s+data\s+analysis\b", re.IGNORECASE),
     "longitudinal data analysis"),
    # bare "panel data" where it implies multi-entity framing
    (re.compile(r"\bpanel\s+data\b", re.IGNORECASE),
     "longitudinal data"),
    # "panel study" / "panel research"
    (re.compile(r"\bpanel\s+(?:study|research)\b", re.IGNORECASE),
     "longitudinal case study"),
    # "population of … banks" → sole population = Union Bank
    (re.compile(
        r"\bpopulation\s+of\s+(?:all\s+)?(?:\d+\s+)?(?:listed\s+)?(?:Nigerian\s+)?(?:commercial\s+)?banks?\b",
        re.IGNORECASE),
     "Union Bank of Nigeria PLC as the sole unit of analysis"),
    # "sampled banks" / "selected banks" / "the banks"
    (re.compile(r"\b(?:sampled|selected)\s+banks?\b", re.IGNORECASE),
     "Union Bank of Nigeria PLC"),
]


def _enforce_ub_scope(text: str) -> str:
    """Replace multi-bank / panel-study framing with Union Bank of Nigeria PLC."""
    for pattern, replacement in _SCOPE_PATTERNS:
        text = pattern.sub(replacement, text)
    return text


def _add_with_union_bank_treatment(doc: Document, text: str) -> None:
    """
    Emit body paragraphs.  Any paragraph that mentions Union Bank is preceded
    by a 'Case Illustration: Union Bank of Nigeria PLC' subsection heading.
    Scope enforcement is applied before the Union Bank check so that converted
    multi-bank references are also treated as Union Bank paragraphs.
    """
    if not text:
        return
    paragraphs = _split_paragraphs(_enforce_ub_scope(str(text)))
    in_ub_block = False
    for para in paragraphs:
        is_ub = bool(re.search(r"\bunion\s+bank\b", para, re.IGNORECASE))
        if is_ub and not in_ub_block:
            _sub_heading(doc, "Case Illustration: Union Bank of Nigeria PLC")
            in_ub_block = True
        elif not is_ub:
            in_ub_block = False
        # _body also calls _enforce_ub_scope; the function is idempotent so
        # double-application is harmless and provides defensive coverage.
        _body(doc, para)


def _add_model_specification(doc: Document, text: str) -> None:
    """
    Render the model specification: the regression equation on its own line
    then each variable definition indented on a separate line.
    """
    if not text:
        return
    text = str(text).strip()
    where_m = re.search(r"\bwhere\b[:\s]", text, re.IGNORECASE)
    if where_m:
        equation = text[: where_m.start()].strip()
        definitions = text[where_m.start():]
        _body(doc, equation)
        for defn in re.split(r"[;\n]", definitions):
            defn = defn.strip()
            if defn:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(3)
                _run(p, defn)
    else:
        # Emit first line as equation, remaining lines indented
        eq_lines = [l.strip() for l in text.split("\n") if l.strip()]
        for i, line in enumerate(eq_lines):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            if i > 0:
                p.paragraph_format.left_indent = Inches(0.5)
            _run(p, line)


def _sort_references(references: list) -> list:
    """Sort APA reference strings alphabetically by the first author surname."""
    def _key(ref: str) -> str:
        m = re.match(r"^([A-Za-z][A-Za-z'\-]*)", str(ref).strip())
        return m.group(1).lower() if m else str(ref).lower()
    return sorted(references, key=_key)


def _parse_ub_table_rows(raw: str) -> list[dict]:
    """
    Parse the union_bank_financial_table string into row dicts.
    Handles pipe-delimited and whitespace-separated formats.
    """
    rows: list[dict] = []
    if not raw:
        return rows

    # Pipe-delimited: | year | npl | car | roa | gc_status | audit_opinion |
    pipe_re = re.compile(
        r"\|\s*(20\d{2})\s*\|"
        r"\s*([^|]+?)\s*\|"
        r"\s*([^|]+?)\s*\|"
        r"\s*([^|]+?)\s*\|"
        r"\s*([^|]+?)\s*\|"
        r"\s*([^|]+?)\s*\|"
    )
    for m in pipe_re.finditer(raw):
        rows.append({
            "year": m.group(1), "npl": m.group(2), "car": m.group(3),
            "roa": m.group(4), "gc_status": m.group(5), "audit_opinion": m.group(6),
        })
    if rows:
        return rows

    # Line-by-line fallback: year followed by tab/comma/multi-space separated fields
    for line in raw.split("\n"):
        line = line.strip()
        year_m = re.match(r"(20\d{2})", line)
        if not year_m:
            continue
        # Normalise mixed delimiters before splitting
        rest = re.sub(r",\s*", "\t", line[year_m.end():].strip())
        fields = [f.strip() for f in re.split(r"\s{2,}|\t", rest) if f.strip()]
        rows.append({
            "year": year_m.group(1),
            "npl": fields[0] if len(fields) > 0 else "",
            "car": fields[1] if len(fields) > 1 else "",
            "roa": fields[2] if len(fields) > 2 else "",
            "gc_status": fields[3] if len(fields) > 3 else "",
            "audit_opinion": fields[4] if len(fields) > 4 else "",
        })
    return rows


def _add_ub_financial_table(doc: Document, raw) -> None:
    """Render the Union Bank financial table as a formatted Word table."""
    raw_str = str(raw) if raw is not None else ""
    rows = _parse_ub_table_rows(raw_str)
    headers = ["Year", "NPL (%)", "CAR (%)", "ROA (%)", "Going Concern Status", "Audit Opinion"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _table_header(table, headers)
    if rows:
        for row in rows:
            _table_row(table, [
                row["year"], row["npl"], row["car"],
                row["roa"], row["gc_status"], row["audit_opinion"],
            ])
    else:
        # Fallback: emit raw as plain text if unparseable
        doc.add_paragraph()
        _body(doc, raw_str)


def _add_ub_narrative_note(doc: Document, raw) -> None:
    """
    Append a plain-prose analytical note derived from the Union Bank financial
    data.  Only figures present in the parsed data are used — nothing is invented.
    """
    raw_str = str(raw) if raw is not None else ""
    by_year = {r["year"]: r for r in _parse_ub_table_rows(raw_str)}

    def _detail(row: dict) -> str:
        parts = []
        for label, key in [("NPL ratio", "npl"), ("CAR", "car"), ("ROA", "roa")]:
            val = row.get(key, "").strip().rstrip("%")
            if val:
                parts.append(f"{label} of {val}%")
        return ", ".join(parts)

    # 2015–2016 modified opinion period
    r15, r16 = by_year.get("2015", {}), by_year.get("2016", {})
    para = (
        "Modified Opinion Period (2015\u20132016): The 2015 and 2016 financial years "
        "represent a period of heightened audit qualification for Union Bank of Nigeria PLC. "
    )
    if r15.get("audit_opinion") or r15.get("gc_status"):
        opinion = r15.get("audit_opinion") or r15.get("gc_status")
        d = _detail(r15)
        para += f"In 2015, the auditors issued a {opinion}" + (f", with {d}" if d else "") + ". "
    if r16.get("audit_opinion") or r16.get("gc_status"):
        opinion = r16.get("audit_opinion") or r16.get("gc_status")
        d = _detail(r16)
        para += f"In 2016, the opinion was {opinion}" + (f", with {d}" if d else "") + ". "
    para += (
        "These consecutive modified opinions are consistent with the asset quality "
        "deterioration documented in CBN Banking Supervision Reports for the same period."
    )
    _body(doc, para)

    # 2018 IFRS 9 transition
    r18 = by_year.get("2018", {})
    para = (
        "IFRS 9 Transition Year (2018): The 2018 financial year marked Union Bank\u2019s "
        "first full reporting period under IFRS 9, which introduced the expected credit loss "
        "(ECL) model requiring forward-looking impairment recognition. "
    )
    d = _detail(r18)
    if d:
        para += f"The bank reported {d} in 2018, reflecting the impact of ECL provisions. "
    if r18.get("audit_opinion") or r18.get("gc_status"):
        opinion = r18.get("audit_opinion") or r18.get("gc_status")
        para += f"The audit opinion for 2018 was: {opinion}. "
    para += (
        "The transition year is therefore a critical inflection point in assessing how "
        "IFRS 9 affected going concern disclosures and reported stability indicators."
    )
    _body(doc, para)

    # 2020 COVID-19
    r20 = by_year.get("2020", {})
    para = (
        "COVID-19 Impact (2020): The 2020 reporting year was materially affected by the "
        "COVID-19 pandemic and the Central Bank of Nigeria\u2019s regulatory forbearance "
        "measures, including loan restructuring guidelines and capital conservation buffers. "
    )
    d = _detail(r20)
    if d:
        para += f"Union Bank\u2019s 2020 indicators showed {d}. "
    if r20.get("audit_opinion") or r20.get("gc_status"):
        opinion = r20.get("audit_opinion") or r20.get("gc_status")
        para += f"The audit opinion noted: {opinion}. "
    para += (
        "The pandemic underscored the relevance of going concern assessments under IAS 1 "
        "during periods of systemic external shock."
    )
    _body(doc, para)

    # 2022 change-in-control
    r22 = by_year.get("2022", {})
    para = (
        "Change-in-Control Event (2022): The 2022 financial year coincided with Union "
        "Bank\u2019s acquisition by Titan Trust Bank Limited, representing a significant "
        "change-in-control event. "
    )
    d = _detail(r22)
    if d:
        para += f"The bank reported {d} for the year. "
    if r22.get("audit_opinion") or r22.get("gc_status"):
        opinion = r22.get("audit_opinion") or r22.get("gc_status")
        para += f"The audit opinion was: {opinion}. "
    para += (
        "This corporate restructuring introduces additional complexity in evaluating "
        "going concern continuity and the forward-looking stability signals embedded "
        "in IFRS-compliant disclosures."
    )
    _body(doc, para)


def _render_gaps(doc: Document, gaps) -> None:
    """Render research gap items with bold 'GAP N:' prefix."""
    if isinstance(gaps, list):
        for gap in gaps:
            gap_str = str(gap).strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            colon_idx = gap_str.find(":") if gap_str.upper().startswith("GAP") else -1
            if colon_idx != -1:
                _run(p, gap_str[: colon_idx + 1], bold=True)
                _run(p, gap_str[colon_idx + 1:])
            else:
                _run(p, gap_str)
    else:
        _body(doc, str(gaps))


def _render_hypotheses(doc: Document, hypotheses) -> None:
    """Render H0N/HNA labelled pairs."""
    if isinstance(hypotheses, list):
        for i, hyp in enumerate(hypotheses, start=1):
            if isinstance(hyp, dict):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                _run(p, f"H0{i}: ", bold=True)
                _run(p, _enforce_ub_scope(str(hyp.get("null", ""))))
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(8)
                _run(p2, f"H{i}A: ", bold=True)
                alt = hyp.get("alternative") or hyp.get("alternate", "")
                _run(p2, _enforce_ub_scope(str(alt)))
            else:
                _body(doc, str(hyp))
    else:
        _body(doc, str(hypotheses))


def _render_numbered_list(doc: Document, items) -> None:
    """Render a numbered list with scope enforcement."""
    if isinstance(items, list):
        for i, item in enumerate(items, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            _run(p, f"{i}. {_enforce_ub_scope(str(item))}")
    else:
        _body(doc, str(items))


def _render_terms(doc: Document, terms) -> None:
    """Render operational definition of terms dict."""
    if isinstance(terms, dict):
        for term, definition in terms.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            _run(p, f"{term}: ", bold=True)
            _run(p, str(definition).strip())
    elif terms:
        _body(doc, str(terms))


def _render_var_table(doc: Document, variables) -> None:
    """Render the variable operationalisation table."""
    if isinstance(variables, list) and variables:
        headers = ["Variable", "Proxy Measure", "Data Source", "Time Period"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _table_header(table, headers)
        for var in variables:
            if isinstance(var, dict):
                _table_row(table, [
                    var.get("variable", ""), var.get("proxy", ""),
                    var.get("source", ""), var.get("period", ""),
                ])
            else:
                _table_row(table, [str(var), "", "", ""])
        doc.add_paragraph()
    else:
        _body(doc, str(variables))


def format_dissertation(proposal: dict) -> Document:
    """
    Transform a raw JSON research proposal into a fully structured,
    publication-ready PhD dissertation as a python-docx Document.
    Consumes the nested chapter_one / chapter_two / chapter_three JSON structure.
    Falls back to old flat-key structure for backward compatibility.
    """
    doc = Document()

    # ── Global margins and default font ───────────────────────────────────────
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    style = doc.styles["Normal"]
    style.font.name = _FONT_NAME
    style.font.size = _BODY_PT

    # ── Nested chapter accessors with flat-key fallback ───────────────────────
    ch1 = proposal.get("chapter_one") or {}
    ch2 = proposal.get("chapter_two") or {}
    ch3 = proposal.get("chapter_three") or {}
    _old_lit = proposal.get("literature_review") or {}
    _old_meth = proposal.get("methodology") or {}

    _bg    = ch1.get("background_to_study") or proposal.get("introduction", "")
    _sop   = ch1.get("statement_of_problem") or proposal.get("statement_of_problem", "")
    _obj   = ch1.get("research_objectives") or proposal.get("research_objectives", [])
    _qs    = ch1.get("research_questions") or proposal.get("research_questions", [])
    _hyp   = ch1.get("research_hypotheses") or proposal.get("research_hypotheses", [])
    _sig   = ch1.get("significance_of_study") or proposal.get("significance_of_study", "")
    _scope = ch1.get("scope_and_delimitation") or proposal.get("scope_and_delimitation", "")
    _terms = ch1.get("operational_definition_of_terms") or {}

    # ── Cover Page ────────────────────────────────────────────────────────────
    cover = proposal.get("cover_page", {})
    if isinstance(cover, dict):
        title      = cover.get("title", "")
        author     = cover.get("author", "")
        institution = cover.get("institution", "")
        department  = cover.get("department", "")
        supervisor  = cover.get("supervisor", "")
        date        = cover.get("date", "")
    else:
        title = str(cover) if cover else ""
        author = institution = department = supervisor = date = ""

    for _ in range(4):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(title_p, title or "PhD Research Proposal", size=_TITLE_PT,
         bold=True, color=_HEADING_COLOR)

    for field in [author, institution, department]:
        if field:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, field)
    if supervisor:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"Supervisor: {supervisor}")
    if date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, date)

    # ── CHAPTER ONE: INTRODUCTION ─────────────────────────────────────────────
    doc.add_page_break()
    _chapter_heading(doc, "CHAPTER ONE \u2014 INTRODUCTION")

    _section_heading(doc, "1.1 Background to the Study")
    _add_with_union_bank_treatment(doc, str(_bg))

    _section_heading(doc, "1.2 Statement of the Problem")
    _add_with_union_bank_treatment(doc, str(_sop))

    _section_heading(doc, "1.3 Research Objectives")
    _render_numbered_list(doc, _obj)

    _section_heading(doc, "1.4 Research Questions")
    _render_numbered_list(doc, _qs)

    _section_heading(doc, "1.5 Research Hypotheses")
    _render_hypotheses(doc, _hyp)

    _section_heading(doc, "1.6 Significance of the Study")
    _body(doc, str(_sig))

    _section_heading(doc, "1.7 Scope and Delimitation of the Study")
    _body(doc, str(_scope))

    _section_heading(doc, "1.8 Operational Definition of Terms")
    _render_terms(doc, _terms)

    # ── CHAPTER TWO: LITERATURE REVIEW ───────────────────────────────────────
    doc.add_page_break()
    _chapter_heading(doc, "CHAPTER TWO \u2014 LITERATURE REVIEW")

    # 2.1 Conceptual Review
    _section_heading(doc, "2.1 Conceptual Review")
    conceptual = ch2.get("conceptual_review") or {}
    if isinstance(conceptual, dict) and conceptual:
        _sub_heading(doc, "2.1.1 IFRS: Concept, Adoption and Framework")
        _body(doc, str(conceptual.get("ifrs_concept_adoption", "")))
        _sub_heading(doc, "2.1.2 Going Concern: Concept and Regulatory Basis")
        _body(doc, str(conceptual.get("going_concern_concept", "")))
        _sub_heading(doc, "2.1.3 Banking Sector Stability: Measures and Determinants")
        _body(doc, str(conceptual.get("banking_stability", "")))
        _sub_heading(doc, "2.1.4 Expected Credit Loss (ECL) under IFRS 9")
        _body(doc, str(conceptual.get("expected_credit_loss_ifrs9", "")))
    else:
        _body(doc, str(conceptual or _old_lit.get("conceptual_review", "")))

    # 2.2 Theoretical Framework
    _section_heading(doc, "2.2 Theoretical Framework")
    theo = ch2.get("theoretical_framework") or {}
    if isinstance(theo, dict) and theo:
        _sub_heading(doc, "2.2.1 Agency Theory")
        _body(doc, str(theo.get("agency_theory", "")))
        _sub_heading(doc, "2.2.2 Signalling Theory")
        _body(doc, str(theo.get("signalling_theory", "")))
        _sub_heading(doc, "2.2.3 Stakeholder Theory")
        _body(doc, str(theo.get("stakeholder_theory", "")))
        _sub_heading(doc, "2.2.4 Positive Accounting Theory")
        _body(doc, str(theo.get("positive_accounting_theory", "")))
    else:
        _body(doc, str(theo
                       or _old_lit.get("theoretical_review", "")
                       or proposal.get("theoretical_framework", "")))

    # 2.3 Empirical Review
    _section_heading(doc, "2.3 Empirical Review")
    empirical = ch2.get("empirical_review") or {}
    if isinstance(empirical, dict) and empirical:
        _sub_heading(doc, "2.3.1 IFRS and Banking Stability \u2014 International Evidence")
        _body(doc, str(empirical.get("ifrs_banking_stability_international", "")))
        _sub_heading(doc, "2.3.2 IFRS and Going Concern \u2014 African Evidence")
        _body(doc, str(empirical.get("ifrs_going_concern_africa", "")))
        _sub_heading(doc, "2.3.3 IFRS 9 ECL and Credit Risk in Nigerian Banks")
        _body(doc, str(empirical.get("ifrs9_ecl_nigerian_banks", "")))
        _sub_heading(doc,
                     "2.3.4 Union Bank of Nigeria PLC \u2014 Prior Studies and Context")
        _body(doc, str(empirical.get("union_bank_prior_studies", "")))
    else:
        _body(doc, str(empirical or _old_lit.get("empirical_review", "")))

    # 2.4 Identification of Research Gaps
    _section_heading(doc, "2.4 Identification of Research Gaps")
    gaps = ch2.get("research_gaps") or proposal.get("research_gaps", [])
    _render_gaps(doc, gaps)

    # 2.5 Summary of Literature Review
    _section_heading(doc, "2.5 Summary of Literature Review")
    _body(doc, str(ch2.get("summary_of_literature", "")))

    # ── CHAPTER THREE: RESEARCH METHODOLOGY ──────────────────────────────────
    doc.add_page_break()
    _chapter_heading(doc, "CHAPTER THREE \u2014 RESEARCH METHODOLOGY")

    _section_heading(doc, "3.1 Research Design")
    _body(doc, str(
        ch3.get("research_design") or _old_meth.get("research_design", "")))

    _section_heading(doc, "3.2 Research Philosophy")
    _body(doc, str(
        ch3.get("research_philosophy") or _old_meth.get("research_philosophy", "")))

    _section_heading(doc, "3.3 Population and Sample")
    _body(doc, str(
        ch3.get("population_and_sample")
        or _old_meth.get("population_and_sampling", "")))

    _section_heading(doc, "3.4 Sources of Data")
    _body(doc, str(
        ch3.get("sources_of_data") or _old_meth.get("data_collection", "")))

    _section_heading(doc, "3.5 Data Collection Procedure")
    _body(doc, str(ch3.get("data_collection_procedure", "")))

    _section_heading(doc, "3.6 Operationalisation of Variables")
    variables = (ch3.get("variable_operationalisation")
                 or proposal.get("data_sources_and_variables", []))
    _render_var_table(doc, variables)

    _section_heading(doc, "3.7 Model Specification")
    _add_model_specification(doc, str(
        ch3.get("model_specification") or _old_meth.get("model_specification", "")))

    _section_heading(doc, "3.8 Method of Data Analysis")
    _body(doc, str(
        ch3.get("method_of_data_analysis")
        or _old_meth.get("analytical_techniques", "")))

    _section_heading(doc, "3.9 Validity and Reliability of Secondary Data")
    _body(doc, str(ch3.get("validity_and_reliability", "")))

    _section_heading(doc, "3.10 Ethical Considerations")
    _body(doc, str(
        ch3.get("ethical_considerations")
        or proposal.get("ethical_considerations", "")))

    # ── REFERENCES ────────────────────────────────────────────────────────────
    doc.add_page_break()
    _chapter_heading(doc, "REFERENCES")
    references = proposal.get("references", [])
    if isinstance(references, list):
        for ref in _sort_references(references):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(str(ref).strip())
            r.font.name = _FONT_NAME
            r.font.size = Pt(11)
    else:
        _body(doc, str(references))

    # ── APPENDIX A ────────────────────────────────────────────────────────────
    doc.add_page_break()
    _chapter_heading(doc,
                     "APPENDIX A \u2014 Union Bank of Nigeria PLC "
                     "Financial Data Table (2015\u20132022)")

    appendices = proposal.get("appendices", {})
    ub_table_raw = ""
    if isinstance(appendices, dict):
        ub_table_raw = appendices.get("union_bank_financial_table", "")
    elif appendices:
        ub_table_raw = str(appendices)

    cap = doc.add_paragraph()
    cap.paragraph_format.space_after = Pt(4)
    _run(cap,
         "Table A1: Union Bank of Nigeria PLC \u2014 "
         "Key Financial and Audit Indicators (2015\u20132022)",
         bold=True)

    _add_ub_financial_table(doc, ub_table_raw)
    doc.add_paragraph()

    _sub_heading(doc, "Appendix Note: Narrative Analysis of Union Bank Transition Events")
    _add_ub_narrative_note(doc, ub_table_raw)

    # ── APPENDIX B ────────────────────────────────────────────────────────────
    doc.add_page_break()
    _chapter_heading(doc,
                     "APPENDIX B \u2014 List of Variables, Proxies and Sources")
    appendix_b = ""
    if isinstance(appendices, dict):
        appendix_b = appendices.get("appendix_b_variables", "")
    _body(doc,
          str(appendix_b) if appendix_b else
          "See Section 3.6 (Operationalisation of Variables) in Chapter Three "
          "for the complete variable operationalisation table.")

    return doc


def main() -> None:
    prompts_path = Path("prompts/synthesis_system.txt")
    if not prompts_path.exists():
        raise FileNotFoundError(f"{prompts_path} not found")

    system_prompt = prompts_path.read_text()

    research = load_research()
    context_str = build_context(research)

    instruction = (
        "Using ALL sources above — [TAV-N:SEARCH] Tavily standard search results, "
        "[TAV-N:EXTRACT] Tavily full-page extractions from authoritative URLs, "
        "[DEEP-N] Tavily deep research synthesis reports, and "
        "the UNION BANK LIVE + BASELINE CASE STUDY DATA — "
        "write a complete PhD research proposal in valid JSON only. "
        "No markdown fences. No preamble. Output only the JSON object. "
        "Follow EVERY instruction in the system prompt exactly. "
        "Use only real, verifiable sources present in the context above for all citations."
    )

    user_message = context_str + "\n\n---\n" + instruction

    tav_count = len(research.get("tavily_results", []))
    deep_count = len(research.get("tavily_research", []))
    print(
        f"[synthesis_agent] Sources loaded: "
        f"Tavily={tav_count}, DeepResearch={deep_count}. "
        f"Context: {len(user_message)} chars. "
        f"Calling OpenRouter ({OPENROUTER_MODEL})..."
    )

    raw = call_openrouter(system_prompt, user_message)

    # Strip accidental markdown code fences
    cleaned = re.sub(r"^```(?:json)?\s*", "", raw.strip(), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned.strip())

    research_dir = Path("research")
    research_dir.mkdir(exist_ok=True)

    try:
        proposal = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        raw_path = research_dir / "proposal_raw.txt"
        raw_path.write_text(raw)
        print(
            f"[synthesis_agent] ERROR: Failed to parse JSON response. "
            f"Raw output saved to {raw_path}. JSONDecodeError: {exc}"
        )
        sys.exit(1)

    output_path = research_dir / "proposal_draft.json"
    output_path.write_text(json.dumps(proposal, indent=2, ensure_ascii=False))

    docx_path = research_dir / "proposal_draft.docx"
    format_dissertation(proposal).save(str(docx_path))

    section_keys = list(proposal.keys())
    print(
        f"[synthesis_agent] Proposal JSON written to {output_path}. "
        f"Formatted dissertation written to {docx_path}. "
        f"Sections found: {section_keys}"
    )


if __name__ == "__main__":
    main()
